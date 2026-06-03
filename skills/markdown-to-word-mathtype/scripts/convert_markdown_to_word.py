#!/usr/bin/env python
from __future__ import annotations

import argparse
import re
import sys
import time
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Any

try:
    import mistune
except ImportError as exc:  # pragma: no cover
    raise SystemExit("缺少 `mistune`，先在 `demo244` 环境里安装对应依赖。") from exc

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor, Emu
except ImportError as exc:  # pragma: no cover
    raise SystemExit("缺少 `python-docx`，先执行 `pip install python-docx`。") from exc


INLINE_MATH_RE = re.compile(r"(?<!\\)\$(.+?)(?<!\\)\$")
INLINE_PLACEHOLDER_RE = re.compile(r"@@IMATH(\d{4})@@")
DISPLAY_TAG_RE = re.compile(r"\\tag\*?\{(.+?)\}\s*$", re.S)
DISPLAY_SUPPRESS_NUMBER_RE = re.compile(r"\\(?:notag|nonumber)\b")
CAPTION_PREFIX_RE = re.compile(r"^[图表]\s*[0-9０-９]+(?:\s*[-.．]\s*[0-9０-９]+)*(?:\s|$)")
MACRO_ENTRY = "MathTypeCommands.UILib.MTCommand_TeXToggle"
STYLE_RGB_BLACK = RGBColor(0x00, 0x00, 0x00)


@dataclass(frozen=True)
class FontSpec:
    east_asia: str
    latin: str
    size_pt: float
    bold: bool = False
    italic: bool = False


@dataclass(frozen=True)
class ParagraphSpec:
    font: FontSpec
    align: int
    space_before_pt: float
    space_after_pt: float
    line_spacing: float
    first_line_indent_pt: float | None = None


@dataclass(frozen=True)
class RunState:
    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False

    def with_updates(self, **kwargs: Any) -> "RunState":
        return replace(self, **kwargs)


@dataclass(frozen=True)
class FormulaJob:
    index: int
    kind: str
    start_marker: str
    end_marker: str


STYLE_SPECS: dict[str, ParagraphSpec] = {
    "Heading 1": ParagraphSpec(
        font=FontSpec("SimHei", "Times New Roman", 16, bold=True),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=12.8,
        space_after_pt=9.6,
        line_spacing=1.5,
        first_line_indent_pt=0,
    ),
    "Heading 2": ParagraphSpec(
        font=FontSpec("SimHei", "Times New Roman", 14, bold=True),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=8.4,
        space_after_pt=5.6,
        line_spacing=1.5,
        first_line_indent_pt=0,
    ),
    "Heading 3": ParagraphSpec(
        font=FontSpec("SimHei", "Times New Roman", 12, bold=True),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=6,
        space_after_pt=3.6,
        line_spacing=1.5,
        first_line_indent_pt=0,
    ),
    "Normal": ParagraphSpec(
        font=FontSpec("SimSun", "Times New Roman", 12),
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before_pt=0,
        space_after_pt=2.4,
        line_spacing=1.5,
        first_line_indent_pt=24,
    ),
    "Markdown Code": ParagraphSpec(
        font=FontSpec("SimSun", "JetBrains Mono", 10.5, bold=True),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=5.25,
        space_after_pt=5.25,
        line_spacing=1.0,
        first_line_indent_pt=0,
    ),
    "Table Paragraph": ParagraphSpec(
        font=FontSpec("SimSun", "Times New Roman", 10.5),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
        first_line_indent_pt=0,
    ),
    "Caption": ParagraphSpec(
        font=FontSpec("SimSun", "Times New Roman", 10.5),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=0,
        space_after_pt=2.4,
        line_spacing=1.5,
        first_line_indent_pt=0,
    ),
    "Display Math": ParagraphSpec(
        font=FontSpec("SimSun", "Times New Roman", 12),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=0,
        space_after_pt=2.4,
        line_spacing=1.5,
        first_line_indent_pt=0,
    ),
}


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="把 Markdown 转成带 MathType 公式的 Word 文档。")
    parser.add_argument("--input", required=True, help="输入 Markdown 文件")
    parser.add_argument("--output", help="输出 Word 文件，默认与输入同名 .docx")
    parser.add_argument(
        "--mathtype-mode",
        choices=("require", "auto", "skip"),
        default="require",
        help="公式转换模式：require=必须转成 MathType，auto=尽量转，skip=跳过",
    )
    parser.add_argument("--overwrite", action="store_true", help="允许覆盖已有输出文件")
    return parser


def ensure_rfonts(r_pr: Any) -> Any:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    return r_fonts


def set_fonts(target: Any, east_asia: str, latin: str) -> None:
    r_pr = target.element.get_or_add_rPr() if hasattr(target, "element") else target._element.get_or_add_rPr()
    r_fonts = ensure_rfonts(r_pr)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)


def apply_style_spec(style: Any, spec: ParagraphSpec) -> None:
    style.font.name = spec.font.latin
    style.font.size = Pt(spec.font.size_pt)
    style.font.bold = spec.font.bold
    style.font.italic = spec.font.italic
    style.font.color.rgb = STYLE_RGB_BLACK
    set_fonts(style, spec.font.east_asia, spec.font.latin)
    fmt = style.paragraph_format
    fmt.alignment = spec.align
    fmt.space_before = Pt(spec.space_before_pt)
    fmt.space_after = Pt(spec.space_after_pt)
    fmt.line_spacing = spec.line_spacing
    if spec.first_line_indent_pt is not None:
        fmt.first_line_indent = Pt(spec.first_line_indent_pt)


def configure_document_styles(doc: Document) -> None:
    styles = doc.styles
    apply_style_spec(styles["Normal"], STYLE_SPECS["Normal"])
    apply_style_spec(styles["Heading 1"], STYLE_SPECS["Heading 1"])
    apply_style_spec(styles["Heading 2"], STYLE_SPECS["Heading 2"])
    apply_style_spec(styles["Heading 3"], STYLE_SPECS["Heading 3"])

    if "Markdown Code" in styles:
        code_style = styles["Markdown Code"]
    else:
        code_style = styles.add_style("Markdown Code", WD_STYLE_TYPE.PARAGRAPH)
    apply_style_spec(code_style, STYLE_SPECS["Markdown Code"])

    if "Caption" in styles:
        caption_style = styles["Caption"]
    else:
        caption_style = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    apply_style_spec(caption_style, STYLE_SPECS["Caption"])

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)


def apply_paragraph_spec(paragraph: Any, spec: ParagraphSpec) -> None:
    fmt = paragraph.paragraph_format
    fmt.alignment = spec.align
    fmt.space_before = Pt(spec.space_before_pt)
    fmt.space_after = Pt(spec.space_after_pt)
    fmt.line_spacing = spec.line_spacing
    if spec.first_line_indent_pt is None:
        fmt.first_line_indent = None
    else:
        fmt.first_line_indent = Pt(spec.first_line_indent_pt)


def apply_run_style(run: Any, font: FontSpec, state: RunState) -> None:
    latin_font = "JetBrains Mono" if state.code else font.latin
    size_pt = 10.5 if state.code else font.size_pt
    run.font.name = latin_font
    run.font.size = Pt(size_pt)
    run.font.bold = font.bold or state.bold or state.code
    run.font.italic = font.italic or state.italic
    run.font.strike = state.strike
    run.font.color.rgb = STYLE_RGB_BLACK
    set_fonts(run, font.east_asia, latin_font)


def extract_block_math(text: str) -> tuple[str, dict[str, str]]:
    lines = text.splitlines()
    result: list[str] = []
    blocks: dict[str, str] = {}
    idx = 0
    block_index = 0
    in_fence = False
    fence_marker = ""

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        fence_match = re.match(r"^(```+|~~~+)", stripped)
        if fence_match:
            marker = fence_match.group(1)[0]
            count = len(fence_match.group(1))
            if not in_fence:
                in_fence = True
                fence_marker = marker * count
            elif stripped.startswith(fence_marker):
                in_fence = False
                fence_marker = ""
            result.append(line)
            idx += 1
            continue

        if in_fence:
            result.append(line)
            idx += 1
            continue

        if stripped.startswith("$$"):
            prefix = line[line.find("$$") + 2 :]
            if stripped.endswith("$$") and stripped != "$$":
                block_index += 1
                key = f"@@BLOCK_MATH_{block_index}@@"
                blocks[key] = stripped[2:-2].strip()
                result.extend(["", key, ""])
                idx += 1
                continue

            math_lines: list[str] = []
            if prefix.strip():
                math_lines.append(prefix)
            idx += 1
            found_end = False
            while idx < len(lines):
                current = lines[idx]
                current_stripped = current.strip()
                if current_stripped.endswith("$$"):
                    suffix_pos = current.rfind("$$")
                    content = current[:suffix_pos]
                    if content.strip():
                        math_lines.append(content)
                    found_end = True
                    idx += 1
                    break
                math_lines.append(current)
                idx += 1
            if not found_end:
                result.append(line)
                result.extend(math_lines)
                continue
            block_index += 1
            key = f"@@BLOCK_MATH_{block_index}@@"
            blocks[key] = "\n".join(math_lines).strip()
            result.extend(["", key, ""])
            continue

        result.append(line)
        idx += 1

    return "\n".join(result), blocks


def extract_inline_math(text: str) -> tuple[str, dict[str, str]]:
    lines = text.splitlines()
    result: list[str] = []
    formulas: dict[str, str] = {}
    in_fence = False
    fence_marker = ""
    index = 0

    for line in lines:
        stripped = line.strip()
        fence_match = re.match(r"^(```+|~~~+)", stripped)
        if fence_match:
            marker = fence_match.group(1)[0]
            count = len(fence_match.group(1))
            if not in_fence:
                in_fence = True
                fence_marker = marker * count
            elif stripped.startswith(fence_marker):
                in_fence = False
                fence_marker = ""
            result.append(line)
            continue

        if in_fence:
            result.append(line)
            continue

        def repl(match: re.Match[str]) -> str:
            nonlocal index
            index += 1
            key = f"@@IMATH{index:04d}@@"
            formulas[key] = match.group(1).strip()
            return key

        result.append(INLINE_MATH_RE.sub(repl, line))

    return "\n".join(result), formulas


def split_display_math_tag(tex: str) -> tuple[str, str | None]:
    normalized = DISPLAY_SUPPRESS_NUMBER_RE.sub("", tex).strip()
    match = DISPLAY_TAG_RE.search(normalized)
    if not match:
        return normalized, None
    tag = match.group(1).strip()
    body = DISPLAY_SUPPRESS_NUMBER_RE.sub("", normalized[: match.start()]).rstrip()
    return body, tag or None


def format_equation_number(tag: str) -> str:
    stripped = tag.strip()
    if not stripped:
        return ""
    if (stripped.startswith("(") and stripped.endswith(")")) or (
        stripped.startswith("[") and stripped.endswith("]")
    ):
        return stripped
    return f"({stripped})"


class MarkdownToWordConverter:
    def __init__(self, base_dir: Path) -> None:
        self.base_dir = base_dir
        self.doc = Document()
        configure_document_styles(self.doc)
        self.md = mistune.create_markdown(renderer="ast", plugins=["table", "strikethrough"])
        self.block_math_map: dict[str, str] = {}
        self.inline_math_map: dict[str, str] = {}
        self.formula_jobs: list[FormulaJob] = []
        self.formula_index = 0

    def next_formula_markers(self, kind: str) -> FormulaJob:
        self.formula_index += 1
        start = f"[[M2W_{self.formula_index:04d}_{kind.upper()}_START]]"
        end = f"[[M2W_{self.formula_index:04d}_{kind.upper()}_END]]"
        job = FormulaJob(self.formula_index, kind, start, end)
        self.formula_jobs.append(job)
        return job

    def convert(self, markdown_text: str) -> tuple[Document, list[FormulaJob]]:
        normalized, self.block_math_map = extract_block_math(markdown_text)
        normalized, self.inline_math_map = extract_inline_math(normalized)
        ast = self.md(normalized)
        for node in ast:
            self.render_block(node)
        return self.doc, self.formula_jobs

    def render_block(self, node: dict[str, Any], list_level: int = 0) -> None:
        handler = getattr(self, f"render_{node['type']}", None)
        if handler is not None:
            handler(node, list_level=list_level)

    def render_heading(self, node: dict[str, Any], list_level: int = 0) -> None:
        level = min(node.get("level", 3), 3)
        paragraph = self.doc.add_paragraph(style=f"Heading {level}")
        self.render_inline_nodes(paragraph, node.get("children", []), STYLE_SPECS[f"Heading {level}"].font, RunState())

    def render_paragraph(self, node: dict[str, Any], list_level: int = 0) -> None:
        placeholder = self.as_block_math_placeholder(node.get("children", []))
        if placeholder is not None:
            self.render_display_math(placeholder)
            return
        image_node = self.as_single_image_node(node.get("children", []))
        if image_node is not None:
            self.render_centered_image(image_node)
            return
        caption_text = self.as_caption_text(node.get("children", []))
        if caption_text is not None:
            self.render_caption(caption_text)
            return
        paragraph = self.doc.add_paragraph(style="Normal")
        if list_level:
            paragraph.paragraph_format.left_indent = Pt(18 * list_level)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        self.render_inline_nodes(paragraph, node.get("children", []), STYLE_SPECS["Normal"].font, RunState())

    def render_block_text(self, node: dict[str, Any], list_level: int = 0) -> None:
        caption_text = self.as_caption_text(node.get("children", []))
        if caption_text is not None:
            self.render_caption(caption_text)
            return
        paragraph = self.doc.add_paragraph(style="Normal")
        if list_level:
            paragraph.paragraph_format.left_indent = Pt(18 * list_level)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        self.render_inline_nodes(paragraph, node.get("children", []), STYLE_SPECS["Normal"].font, RunState())

    def render_block_code(self, node: dict[str, Any], list_level: int = 0) -> None:
        paragraph = self.doc.add_paragraph(style="Markdown Code")
        run = paragraph.add_run(node.get("text", "").rstrip("\n"))
        apply_run_style(run, STYLE_SPECS["Markdown Code"].font, RunState(code=True))

    def render_list(self, node: dict[str, Any], list_level: int = 0) -> None:
        ordered = node.get("ordered", False)
        for idx, item in enumerate(node.get("children", []), start=1):
            prefix = f"{idx}. " if ordered else "- "
            paragraph = self.doc.add_paragraph(style="Normal")
            paragraph.paragraph_format.left_indent = Pt(18 * list_level)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            prefix_run = paragraph.add_run(prefix)
            apply_run_style(prefix_run, STYLE_SPECS["Normal"].font, RunState())
            first_text_block = True
            for child in item.get("children", []):
                child_type = child["type"]
                if child_type in {"block_text", "paragraph"}:
                    if first_text_block:
                        self.render_inline_nodes(paragraph, child.get("children", []), STYLE_SPECS["Normal"].font, RunState())
                        first_text_block = False
                    else:
                        self.render_block(child, list_level=list_level + 1)
                elif child_type == "list":
                    self.render_list(child, list_level=list_level + 1)
                else:
                    self.render_block(child, list_level=list_level + 1)

    def render_block_quote(self, node: dict[str, Any], list_level: int = 0) -> None:
        for child in node.get("children", []):
            paragraph = self.doc.add_paragraph(style="Normal")
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            prefix = paragraph.add_run("引文：")
            apply_run_style(prefix, STYLE_SPECS["Normal"].font, RunState())
            if child["type"] in {"paragraph", "block_text"}:
                self.render_inline_nodes(paragraph, child.get("children", []), STYLE_SPECS["Normal"].font, RunState())
            else:
                self.render_block(child, list_level=list_level)

    def render_table(self, node: dict[str, Any], list_level: int = 0) -> None:
        rows: list[list[dict[str, Any]]] = []
        for group in node.get("children", []):
            if group["type"] == "table_head":
                rows.append(group.get("children", []))
            elif group["type"] == "table_body":
                for row in group.get("children", []):
                    rows.append(row.get("children", []))
        if not rows:
            return
        column_count = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(rows):
            for col_idx, cell_node in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                apply_paragraph_spec(paragraph, STYLE_SPECS["Table Paragraph"])
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.render_inline_nodes(paragraph, cell_node.get("children", []), STYLE_SPECS["Table Paragraph"].font, RunState())

    def render_thematic_break(self, node: dict[str, Any], list_level: int = 0) -> None:
        paragraph = self.doc.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run("------------------------------------------------------------")
        apply_run_style(run, STYLE_SPECS["Normal"].font, RunState())

    def render_newline(self, node: dict[str, Any], list_level: int = 0) -> None:
        return

    def as_block_math_placeholder(self, children: list[dict[str, Any]]) -> str | None:
        if len(children) != 1 or children[0]["type"] != "text":
            return None
        text = children[0].get("text", "").strip()
        return text if text in self.block_math_map else None

    def render_display_math(self, placeholder: str) -> None:
        tex = " ".join(line.strip() for line in self.block_math_map[placeholder].splitlines() if line.strip())
        tex, tag = split_display_math_tag(tex)
        job = self.next_formula_markers("block")

        if tag:
            paragraph = self.doc.add_paragraph()
            apply_paragraph_spec(paragraph, STYLE_SPECS["Display Math"])
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.configure_numbered_display_tabs(paragraph)
            display_tex = tex if tex.lstrip().startswith(r"\displaystyle") else rf"\displaystyle {tex}"

            tab_run = paragraph.add_run("\t")
            apply_run_style(tab_run, STYLE_SPECS["Display Math"].font, RunState())

            formula_run = paragraph.add_run(f"{job.start_marker}${display_tex}${job.end_marker}")
            apply_run_style(formula_run, STYLE_SPECS["Display Math"].font, RunState())

            number_tab_run = paragraph.add_run("\t")
            apply_run_style(number_tab_run, STYLE_SPECS["Display Math"].font, RunState())

            number_run = paragraph.add_run(format_equation_number(tag))
            apply_run_style(number_run, STYLE_SPECS["Display Math"].font, RunState())
            return

        paragraph = self.doc.add_paragraph()
        apply_paragraph_spec(paragraph, STYLE_SPECS["Display Math"])
        run = paragraph.add_run(f"{job.start_marker}\\[{tex}\\]{job.end_marker}")
        apply_run_style(run, STYLE_SPECS["Display Math"].font, RunState())

    def configure_numbered_display_tabs(self, paragraph: Any) -> None:
        tab_stops = paragraph.paragraph_format.tab_stops
        for idx in range(len(tab_stops), 0, -1):
            tab_stops[idx - 1].clear()

        section = self.doc.sections[-1]
        text_width = section.page_width - section.left_margin - section.right_margin
        center_pos = Emu(text_width // 2)
        right_pos = Emu(text_width)
        tab_stops.add_tab_stop(center_pos, WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES)
        tab_stops.add_tab_stop(right_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    def as_single_image_node(self, children: list[dict[str, Any]]) -> dict[str, Any] | None:
        if len(children) == 1 and children[0]["type"] == "image":
            return children[0]
        return None

    def as_caption_text(self, children: list[dict[str, Any]]) -> str | None:
        if not children or any(node["type"] not in {"strong", "emphasis"} for node in children):
            return None
        text = self.flatten_inline_text(children)
        if text is None:
            return None
        normalized = re.sub(r"\s+", " ", text).strip()
        if not normalized or not CAPTION_PREFIX_RE.match(normalized):
            return None
        return normalized

    def flatten_inline_text(self, nodes: list[dict[str, Any]]) -> str | None:
        parts: list[str] = []

        def walk(node: dict[str, Any]) -> bool:
            node_type = node["type"]
            if node_type == "text":
                parts.append(node.get("text", ""))
                return True
            if node_type in {"strong", "emphasis", "strikethrough"}:
                for child in node.get("children", []):
                    if not walk(child):
                        return False
                return True
            return False

        for node in nodes:
            if not walk(node):
                return None
        return "".join(parts)

    def render_caption(self, text: str) -> None:
        paragraph = self.doc.add_paragraph(style="Caption")
        apply_paragraph_spec(paragraph, STYLE_SPECS["Caption"])
        run = paragraph.add_run(text)
        apply_run_style(run, STYLE_SPECS["Caption"].font, RunState())

    def render_centered_image(self, node: dict[str, Any]) -> None:
        paragraph = self.doc.add_paragraph(style="Normal")
        apply_paragraph_spec(paragraph, STYLE_SPECS["Normal"])
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        src = node.get("src", "")
        if src and not re.match(r"^[a-zA-Z]+://", src):
            candidate = (self.base_dir / src).resolve()
            if candidate.exists():
                paragraph.add_run().add_picture(str(candidate), width=Cm(14))
                return
        alt = node.get("alt", "image")
        run = paragraph.add_run(f"[图片: {alt}]")
        apply_run_style(run, STYLE_SPECS["Caption"].font, RunState())

    def render_inline_nodes(self, paragraph: Any, nodes: list[dict[str, Any]], base_font: FontSpec, state: RunState) -> None:
        nodes = self.merge_adjacent_text_nodes(nodes)
        for node in nodes:
            node_type = node["type"]
            if node_type == "text":
                self.add_text_with_inline_math(paragraph, node.get("text", ""), base_font, state)
            elif node_type == "strong":
                self.render_inline_nodes(paragraph, node.get("children", []), base_font, state.with_updates(bold=True))
            elif node_type == "emphasis":
                self.render_inline_nodes(paragraph, node.get("children", []), base_font, state.with_updates(italic=True))
            elif node_type == "strikethrough":
                self.render_inline_nodes(paragraph, node.get("children", []), base_font, state.with_updates(strike=True))
            elif node_type == "codespan":
                run = paragraph.add_run(node.get("text", ""))
                apply_run_style(run, base_font, state.with_updates(code=True))
            elif node_type in {"softbreak", "linebreak"}:
                paragraph.add_run("\n")
            elif node_type == "link":
                self.render_inline_nodes(paragraph, node.get("children", []), base_font, state)
                run = paragraph.add_run(f" ({node.get('link', '')})")
                apply_run_style(run, base_font, state)
            elif node_type == "image":
                self.render_image(paragraph, node, base_font, state)
            else:
                self.render_inline_nodes(paragraph, node.get("children", []), base_font, state)

    def merge_adjacent_text_nodes(self, nodes: list[dict[str, Any]]) -> list[dict[str, Any]]:
        merged: list[dict[str, Any]] = []
        buffer: list[str] = []

        def flush() -> None:
            if buffer:
                merged.append({"type": "text", "text": "".join(buffer)})
                buffer.clear()

        for node in nodes:
            if node["type"] == "text":
                buffer.append(node.get("text", ""))
            else:
                flush()
                merged.append(node)
        flush()
        return merged

    def add_text_with_inline_math(self, paragraph: Any, text: str, base_font: FontSpec, state: RunState) -> None:
        cursor = 0
        for placeholder in INLINE_PLACEHOLDER_RE.finditer(text):
            raw_segment = text[cursor : placeholder.start()]
            self.add_raw_inline_math_runs(paragraph, raw_segment, base_font, state)
            key = placeholder.group(0)
            formula = self.inline_math_map.get(key)
            if formula is None:
                self.add_plain_run(paragraph, key, base_font, state)
            else:
                job = self.next_formula_markers("inline")
                self.add_plain_run(paragraph, f"{job.start_marker}${formula}${job.end_marker}", base_font, state)
            cursor = placeholder.end()
        self.add_raw_inline_math_runs(paragraph, text[cursor:], base_font, state)

    def add_raw_inline_math_runs(self, paragraph: Any, text: str, base_font: FontSpec, state: RunState) -> None:
        cursor = 0
        for match in INLINE_MATH_RE.finditer(text):
            prefix = text[cursor : match.start()]
            self.add_plain_run(paragraph, prefix.replace("\\$", "$"), base_font, state)
            formula = match.group(1).strip()
            job = self.next_formula_markers("inline")
            self.add_plain_run(paragraph, f"{job.start_marker}${formula}${job.end_marker}", base_font, state)
            cursor = match.end()
        self.add_plain_run(paragraph, text[cursor:].replace("\\$", "$"), base_font, state)

    def add_plain_run(self, paragraph: Any, text: str, base_font: FontSpec, state: RunState) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        apply_run_style(run, base_font, state)

    def render_image(self, paragraph: Any, node: dict[str, Any], base_font: FontSpec, state: RunState) -> None:
        src = node.get("src", "")
        if src and not re.match(r"^[a-zA-Z]+://", src):
            candidate = (self.base_dir / src).resolve()
            if candidate.exists():
                paragraph.add_run().add_picture(str(candidate), width=Cm(14))
                return
        alt = node.get("alt", "image")
        run = paragraph.add_run(f"[图片: {alt}]")
        apply_run_style(run, base_font, state)


def find_text_range(doc: Any, target: str, start: int = 0) -> Any | None:
    rng = doc.Range(Start=start, End=doc.Content.End)
    finder = rng.Find
    finder.ClearFormatting()
    finder.Text = target
    if not finder.Execute():
        return None
    return doc.Range(Start=rng.Start, End=rng.End)


def candidate_addins() -> list[Path]:
    root = Path(r"C:\Program Files (x86)\MathType\Office Support")
    result: list[Path] = []
    for arch in ("64", "32"):
        folder = root / arch
        if not folder.exists():
            continue
        for year in ("2016", "2013", "2010"):
            path = folder / f"MathType Commands {year}.dotm"
            if path.exists():
                result.append(path)
    return result


def resolve_macro_name(word: Any) -> str | None:
    for idx in range(1, word.Templates.Count + 1):
        template = word.Templates(idx)
        if template.Name.startswith("MathType Commands"):
            return f"'{template.Name}'!{MACRO_ENTRY}"
    for path in candidate_addins():
        try:
            addin = word.AddIns.Add(str(path), True)
            addin.Installed = True
        except Exception:
            continue
    for idx in range(1, word.Templates.Count + 1):
        template = word.Templates(idx)
        if template.Name.startswith("MathType Commands"):
            return f"'{template.Name}'!{MACRO_ENTRY}"
    return None


def cleanup_markers(doc: Any, jobs: list[FormulaJob]) -> None:
    for job in jobs:
        for marker in (job.end_marker, job.start_marker):
            rng = find_text_range(doc, marker)
            if rng is not None:
                rng.Text = ""


def run_mathtype_macro(word: Any, macro_name: str, retries: int = 3) -> None:
    last_error: Exception | None = None
    for attempt in range(retries):
        try:
            word.Run(macro_name)
            return
        except Exception as exc:
            last_error = exc
            time.sleep(0.35 * (attempt + 1))
    if last_error is not None:
        raise last_error


def convert_formulas_with_mathtype(output_path: Path, jobs: list[FormulaJob], mode: str) -> bool:
    if not jobs or mode == "skip":
        return False

    try:
        import win32com.client
    except ImportError:
        if mode == "require":
            raise RuntimeError("缺少 `pywin32`，无法调用 Word + MathType。")
        return False

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = word.Documents.Open(str(output_path))
    temp_output = output_path.with_name(f"{output_path.stem}.__mathtype__.docx")
    if temp_output.exists():
        temp_output.unlink()
    saved_temp = False

    def save_to_temp() -> None:
        nonlocal saved_temp
        doc.SaveAs2(str(temp_output), FileFormat=16)
        saved_temp = True

    try:
        macro_name = resolve_macro_name(word)
        if not macro_name:
            cleanup_markers(doc, jobs)
            save_to_temp()
            if mode == "require":
                raise RuntimeError("没有找到可用的 MathType Word 模板，公式未转换。")
            return False

        for job in jobs:
            start_range = find_text_range(doc, job.start_marker)
            if start_range is None:
                continue
            end_range = find_text_range(doc, job.end_marker, start=start_range.End)
            if end_range is None:
                continue
            inner = doc.Range(Start=start_range.End, End=end_range.Start)
            if not inner.Text.strip():
                continue
            inner.Select()
            run_mathtype_macro(word, macro_name)
            time.sleep(0.15)
            end_range = find_text_range(doc, job.end_marker)
            start_range = find_text_range(doc, job.start_marker)
            if end_range is not None:
                end_range.Text = ""
            if start_range is not None:
                start_range.Text = ""

        save_to_temp()
        return True
    except Exception:
        cleanup_markers(doc, jobs)
        save_to_temp()
        if mode == "require":
            raise
        return False
    finally:
        doc.Close(False)
        word.Quit()
        if saved_temp:
            temp_output.replace(output_path)


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        raise SystemExit(f"输入文件不存在：{input_path}")

    output_path = Path(args.output).expanduser().resolve() if args.output else input_path.with_suffix(".docx")
    if output_path.exists() and not args.overwrite:
        raise SystemExit(f"输出文件已存在：{output_path}；如需覆盖，加 `--overwrite`。")

    markdown_text = input_path.read_text(encoding="utf-8").lstrip("\ufeff")
    converter = MarkdownToWordConverter(base_dir=input_path.parent)
    document, jobs = converter.convert(markdown_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    converted = convert_formulas_with_mathtype(output_path, jobs, args.mathtype_mode)
    if converted:
        print(f"已生成并完成 MathType 转换：{output_path}")
    elif jobs and args.mathtype_mode != "skip":
        print(f"已生成 Word，但公式保留为 TeX 文本：{output_path}")
    else:
        print(f"已生成 Word：{output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
